import os
import re
import logging
from docx import Document
import pandas as pd
from typing import List, Optional
from datetime import datetime
from copy import deepcopy
from num2words import num2words

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# Add this at the top of your data_mapper.py file, right after the imports
COLUMN_MAPPING = {
    'invoicenumber': 'INVOICE_NUMBER',
    'invoicedate': 'INVOICE_DATE',
    'recipientgstin': 'RECIPIENT_GSTIN',
    'recipientname': 'RECIPIENT_NAME',
    'recipientaddress': 'RECIPIENT_ADDRESS',
    'recipientstate': 'RECIPIENT_STATE',
    'recipientpincode': 'RECIPIENT_PINCODE',
    'recipientnumber': 'RECIPIENT_NUMBER',
    'suppliergstin': 'SUPPLIER_GSTIN',
    'suppliername': 'SUPPLIER_NAME',
    'supplieraddress': 'SUPPLIER_ADDRESS',
    'supplierstate': 'SUPPLIER__STATE',  # Note double underscore
    'supplierpincode': 'SUPPLIER_PINCODE',
    'suppliernumber': 'SUPPLIER_NUMBER',
    'amount': 'AMOUNT',
    'regoffice': 'REG_OFFICE',
    'cin': 'CIN',
    'e-mail': 'E_MAIL',
    'website': 'WEBSITE',
    'cgst': 'CGST',
    'sgst': 'SGST',
    'utgst': 'UTGST',
    'igst': 'IGST'
}

def normalize_column_names(df: pd.DataFrame) -> pd.DataFrame:
    """More flexible column name normalization"""
    df.columns = [
        col.strip()
        .replace(' ', '_')
        .replace('-', '_')
        .upper()
        for col in df.columns
    ]
    return df


def map_data_to_docx(template_path: str, data: pd.DataFrame, output_folder: str) -> Optional[List[str]]:
    """Generate one DOCX per data row with all placeholders filled"""
    try:
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        if data.empty:
            raise ValueError("No data provided")

        os.makedirs(output_folder, exist_ok=True)
        generated_files = []
        template_doc = Document(template_path)

        # Pre-process all placeholders in the template
        template_placeholders = scan_template_placeholders(template_path)
        logging.info(f"Found placeholders in template: {template_placeholders}")

        # DEBUG: Print template placeholders and data columns
        print("\n=== DEBUG INFORMATION ===")
        print("Template placeholders:", template_placeholders)
        print("Data columns:", data.columns.tolist())
        print("First row data:", dict(data.iloc[0]))
        print("=======================\n")
        # After loading your data
        print("Normalized columns:", [col.lower().replace('_', '') for col in data.columns])
        print("Normalized placeholders:", [ph.lower().replace(' ', '') for ph in template_placeholders])

        for idx, row in data.iterrows():
            try:
                doc = deepcopy(template_doc)
                row_data = prepare_row_data(row, template_placeholders)

                # DEBUG: Print matched data for each row
                print(f"\nProcessing row {idx}:")
                print("Row data keys:", row.keys())
                print("Matched data:", row_data)

                # Right after scanning template placeholders
                print("\n=== PLACEHOLDER MAPPING ===")
                for ph in template_placeholders:
                    norm_ph = ph.lower().replace(' ', '').replace('.', '').replace('-', '')
                    data_key = COLUMN_MAPPING.get(norm_ph, "NO MATCH")
                    print(f"Template: {ph:25} → Data: {data_key}")
                print("==========================\n")

                # Process entire document structure
                replace_all_placeholders(doc, row_data)

                # Save output
                output_path = generate_output_path(output_folder, row_data, idx)
                doc.save(output_path)
                generated_files.append(output_path)
                logging.info(f"Generated: {output_path}")

            except Exception as e:
                logging.error(f"Row {idx + 1} error: {str(e)}")
                continue

        return generated_files if generated_files else None

    except Exception as e:
        logging.error(f"Fatal error: {str(e)}")
        return None

def scan_template_placeholders(template_path: str) -> set:
    """Find all unique placeholders in the template"""
    doc = Document(template_path)
    placeholders = set()

    def scan_text(text: str):
        return {m.group(1) for m in re.finditer(r'\{\{(.*?)\}\}', text)}

    # Scan paragraphs
    for para in doc.paragraphs:
        placeholders.update(scan_text(para.text))
        for run in para.runs:
            placeholders.update(scan_text(run.text))

    # Scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    placeholders.update(scan_text(para.text))
                if cell.tables:  # Nested tables
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for para in nested_cell.paragraphs:
                                    placeholders.update(scan_text(para.text))

    return placeholders


def prepare_row_data(row: pd.Series, template_placeholders: set) -> dict:
    """Match data columns to template placeholders with exact mapping"""
    row_data = {}

    # Process each template placeholder
    for ph in template_placeholders:
        # Normalize the placeholder name
        norm_ph = ph.lower().replace(' ', '').replace('.', '').replace('-', '')

        # Special handling for amount_in_words
        if norm_ph == 'amountinwords':
            try:
                amount = float(row['AMOUNT'])
                row_data[ph] = num2words(amount, lang='en_IN').title() + " Rupees Only"
            except:
                row_data[ph] = ""
            continue

        # Find matching column using our mapping
        data_key = COLUMN_MAPPING.get(norm_ph)
        if data_key and data_key in row:
            value = row[data_key]
            # Convert numpy types to native Python
            if hasattr(value, 'item'):
                value = value.item()
            row_data[ph] = format_value(value, ph)
        else:
            row_data[ph] = ""
            logging.warning(f"No data mapping for placeholder: {ph}")

    return row_data

def replace_all_placeholders(doc, row_data):
    """Replace placeholders throughout entire document"""
    # Process paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, row_data)

    # Process tables
    for table in doc.tables:
        process_table(table, row_data)

    # Process headers and footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, row_data)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, row_data)


def process_table(table, row_data):
    """Handle complex table structures with merged cells"""
    for row in table.rows:
        for cell in row.cells:
            # Join all paragraphs in cell to handle split placeholders
            full_text = '\n'.join(p.text for p in cell.paragraphs)

            # Replace placeholders
            for ph, value in row_data.items():
                full_text = full_text.replace(f'{{{{{ph}}}}}', str(value))

            # Clear and rebuild cell content
            cell.text = ''
            cell.add_paragraph(full_text)

        if len(table.rows) >= 7:  # Your recipient/supplier table
            # Left column (recipient)
            table.cell(1, 0).text = str(row_data.get('Recipient Name', ''))
            table.cell(2, 0).text = str(row_data.get('Recipient Number', ''))
            table.cell(3, 0).text = str(row_data.get('Recipient Address', ''))
            table.cell(4, 0).text = str(row_data.get('Recipient Pincode', ''))
            table.cell(5, 0).text = str(row_data.get('Recipient State', ''))
            table.cell(6, 0).text = str(row_data.get('Recipient State Code', ''))
            table.cell(7, 0).text = str(row_data.get('Recipient GSTIN', ''))
            # ... etc for other recipient fields

            # Right column (supplier)
            table.cell(1, 1).text = str(row_data.get('Supplier Name', ''))
            table.cell(2, 1).text = str(row_data.get('Supplier Number', ''))
            table.cell(3, 1).text = str(row_data.get('Supplier Address', ''))
            table.cell(4, 1).text = str(row_data.get('Supplier Pincode', ''))
            table.cell(5, 1).text = str(row_data.get('Supplier State', ''))
            table.cell(6, 1).text = str(row_data.get('Supplier State Code', ''))
            table.cell(7, 1).text = str(row_data.get('Supplier GSTIN', ''))
            # ... etc for other supplier fields


def replace_in_paragraph(paragraph, row_data):
    """Replace placeholders while preserving formatting"""
    # First check if paragraph has any placeholders
    if not any('{{' in run.text for run in paragraph.runs):
        return

    # Store original runs with formatting
    original_runs = []
    for run in paragraph.runs:
        original_runs.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font': run.font.name,
            'size': run.font.size
        })

    # Build full paragraph text
    full_text = ''.join(r['text'] for r in original_runs)

    # Perform all replacements
    modified = False
    for ph, value in row_data.items():
        if f'{{{{{ph}}}}}' in full_text:
            full_text = full_text.replace(f'{{{{{ph}}}}}', str(value))
            modified = True

    if not modified:
        return

    # Rebuild paragraph with original formatting
    paragraph.clear()
    current_pos = 0

    for run in original_runs:
        run_text = full_text[current_pos:current_pos + len(run['text'])]
        if run_text:  # Only add run if there's content
            new_run = paragraph.add_run(run_text)
            # Preserve formatting
            new_run.bold = run['bold']
            new_run.italic = run['italic']
            new_run.underline = run['underline']
            if run['font']:
                new_run.font.name = run['font']
            if run['size']:
                new_run.font.size = run['size']

        current_pos += len(run['text'])


def format_value(value, key=None):
    """Handle special formatting and numpy types"""
    if pd.isna(value):
        return ""

    # Convert numpy types to native Python types
    if hasattr(value, 'item'):
        value = value.item()

    # Special formatting for Amount
    if key and key.lower() == 'amount':
        try:
            return "{:,.2f}".format(float(value))
        except:
            return str(value)

    # Rest of your formatting logic...
    return str(value)

def generate_output_path(output_folder: str, row_data: dict, idx: int) -> str:
    """Generate output path with invoice number if available"""
    invoice_num = str(row_data.get('INVOICE_NUMBER', idx + 1)).strip().replace('/', '-')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(output_folder, f"ISD_Invoice_{invoice_num}_{timestamp}.docx")

