import os
import re
import logging
from docx import Document
import pandas as pd
from docx.shared import Pt
from typing import Dict, List, Optional, Set
from num2words import num2words
from PyQt5.QtWidgets import QMessageBox
from datetime import datetime

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

class DocxFiller:
    def __init__(self, parent=None):
        self.parent = parent
        self.default_font_size = 10  # Default font size in points

    def fill_template(self, template_path: str, data: pd.DataFrame, output_folder: str) -> Optional[List[str]]:
        """Main function to fill DOCX templates with data"""
        try:
            if not self.validate_inputs(template_path, data, output_folder):
                return None

            os.makedirs(output_folder, exist_ok=True)
            generated_files = []
            template_placeholders = self.extract_placeholders(template_path)

            # Check for missing placeholders in data
            missing_placeholders = self.check_missing_placeholders(template_placeholders, data)
            if missing_placeholders:
                QMessageBox.warning(
                    self.parent,
                    "Warning",
                    f"Template placeholders not found in data:\n{', '.join(missing_placeholders)}"
                )

            for idx, row in data.iterrows():
                try:
                    doc = Document(template_path)
                    row_data = self.prepare_row_data(row, template_placeholders)

                    if not self.replace_placeholders_in_document(doc, row_data):
                        continue

                    output_path = self.generate_output_path(output_folder, row_data, idx)
                    doc.save(output_path)
                    generated_files.append(output_path)
                    logging.info(f"Generated: {output_path}")

                except Exception as e:
                    logging.error(f"Error processing row {idx + 1}: {str(e)}", exc_info=True)
                    continue

            return generated_files if generated_files else None

        except Exception as e:
            logging.error(f"Fatal error: {str(e)}", exc_info=True)
            QMessageBox.critical(self.parent, "Error", f"Failed to fill templates: {str(e)}")
            return None

    def validate_inputs(self, template_path: str, data: pd.DataFrame, output_folder: str) -> bool:
        """Validate all input parameters"""
        if not os.path.exists(template_path):
            QMessageBox.critical(self.parent, "Error", f"Template file not found: {template_path}")
            return False

        if data.empty:
            QMessageBox.critical(self.parent, "Error", "No data provided in DataFrame")
            return False

        try:
            os.makedirs(output_folder, exist_ok=True)
            return True
        except Exception as e:
            QMessageBox.critical(self.parent, "Error", f"Output folder not writable: {str(e)}")
            return False

    def check_missing_placeholders(self, template_placeholders: Set[str], data: pd.DataFrame) -> Set[str]:
        """Check which template placeholders are missing from the data columns"""
        data_columns = {col.lower() for col in data.columns}
        return {
            ph for ph in template_placeholders
            if ph.lower() not in data_columns
        }

    def prepare_row_data(self, row: pd.Series, template_placeholders: Set[str]) -> Dict[str, str]:
        """Prepare complete row data with all required fields and proper formatting"""
        row_data = {}

        # First handle special fields
        if 'amount_in_words' in template_placeholders and 'AMOUNT' in row:
            try:
                amount = float(row['AMOUNT'])
                words = num2words(amount, lang='en_IN').title()
                words = words.replace(' And ', ' and ')  # Fix capitalization
                row_data['amount_in_words'] = f"{words} Rupees Only"
            except Exception as e:
                logging.error(f"Amount to words failed: {str(e)}")
                row_data['amount_in_words'] = ""

        # Process all placeholders in template
        for ph in template_placeholders:
            if ph == 'amount_in_words':
                continue  # Already handled

            # Try to find matching column with flexible matching
            matched_col = None
            norm_ph = ph.lower().replace(' ', '_').replace('-', '_')

            for col in row.index:
                norm_col = col.lower().replace(' ', '_').replace('-', '_')
                if norm_col == norm_ph:
                    matched_col = col
                    break

            if matched_col:
                value = row[matched_col]
                # Convert numpy types to native Python
                if hasattr(value, 'item'):
                    value = value.item()
                row_data[ph] = self.format_value(value, ph)
            else:
                row_data[ph] = ""
                logging.warning(f"No data mapping for placeholder: {ph}")

        return row_data

    def replace_placeholders_in_document(self, doc: Document, row_data: Dict[str, str]) -> bool:
        """Replace placeholders throughout document with formatting preservation"""
        try:
            # Process all paragraphs in main document
            for paragraph in doc.paragraphs:
                self.replace_in_paragraph(paragraph, row_data)

            # Process all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_in_paragraph(paragraph, row_data)

            # Process headers and footers
            for section in doc.sections:
                for header in [section.header, section.first_page_header]:
                    if header:
                        for paragraph in header.paragraphs:
                            self.replace_in_paragraph(paragraph, row_data)

                for footer in [section.footer, section.first_page_footer]:
                    if footer:
                        for paragraph in footer.paragraphs:
                            self.replace_in_paragraph(paragraph, row_data)

            return True

        except Exception as e:
            logging.error(f"Error replacing placeholders: {str(e)}", exc_info=True)
            return False

    def replace_in_paragraph(self, paragraph, row_data: Dict[str, str]):
        """Replace placeholders in a paragraph while preserving formatting"""
        # First combine all runs to handle split placeholders
        full_text = ''.join(run.text for run in paragraph.runs)

        # Skip if no placeholders
        if not any(f'{{{{{ph}}}}}' in full_text for ph in row_data):
            return

        # Perform all replacements
        modified_text = full_text
        for ph, value in row_data.items():
            modified_text = modified_text.replace(f'{{{{{ph}}}}}', value)

        # Only update if changes were made
        if modified_text != full_text:
            # Clear existing content
            paragraph.clear()

            # Add new content with preserved formatting
            run = paragraph.add_run(modified_text)
            run.font.size = Pt(self.default_font_size)

            # Preserve other formatting from first run if available
            if paragraph.runs and paragraph.runs[0].font.name:
                run.font.name = paragraph.runs[0].font.name

    def format_value(self, value, key=None) -> str:
        """Format values with special handling for certain fields"""
        if pd.isna(value):
            return ""

        # Handle numpy types
        if hasattr(value, 'item'):
            value = value.item()

        # Special formatting for amounts
        if key and 'amount' in key.lower() and isinstance(value, (int, float)):
            return "{:,.2f}".format(value)

        # Special handling for GSTIN (format with spaces)
        if key and 'gstin' in key.lower() and isinstance(value, str) and len(value) == 15:
            return f"{value[:2]} {value[2:5]} {value[5:7]} {value[7:12]} {value[12:15]}"

        return str(value).strip()

    def extract_placeholders(self, template_path: str) -> Set[str]:
        """Extract all unique placeholders from a DOCX template"""
        doc = Document(template_path)
        placeholders = set()
        placeholder_pattern = re.compile(r'\{\{\s*(.*?)\s*\}\}')  # Handles whitespace

        def extract_from_text(text: str):
            """Inner function to extract placeholders from text"""
            return {match.strip() for match in placeholder_pattern.findall(text)}

        # Process all document components
        components = [
            doc.paragraphs,
            *[cell.paragraphs for table in doc.tables
              for row in table.rows
              for cell in row.cells],
            *[section.header.paragraphs for section in doc.sections],
            *[section.footer.paragraphs for section in doc.sections]
        ]

        for paragraphs in components:
            for paragraph in paragraphs:
                placeholders.update(extract_from_text(paragraph.text))
                for run in paragraph.runs:
                    placeholders.update(extract_from_text(run.text))

        return {ph for ph in placeholders if ph}  # Remove empty strings

    def generate_output_path(self, output_folder: str, row_data: dict, idx: int) -> str:
        """Generate output path with invoice number if available"""
        invoice_num = str(row_data.get('INVOICE_NUMBER', idx + 1)).strip()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return os.path.join(output_folder, f"Document_{invoice_num}_{timestamp}.docx")