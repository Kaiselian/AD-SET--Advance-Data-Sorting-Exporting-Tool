import sys
import os
import pandas as pd
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QFileDialog, QMessageBox, QLabel, QTableWidget, QTableWidgetItem, QLineEdit, QComboBox,
    QDialog, QListWidget, QListWidgetItem, QFormLayout, QDialogButtonBox, QScrollArea, QGraphicsView, QGraphicsScene, QGraphicsRectItem
)
from PyQt5.QtCore import Qt, QFileInfo, QStandardPaths
from PyQt5.QtGui import QPixmap
from fpdf import FPDF
import pdfplumber
import fitz
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from docx import Document
import json
import logging


# Set up logging
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


class MainApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Advanced Data Search & Export Tool 2.2")
        self.setGeometry(100, 100, 1200, 800)

        # Initialize variables
        self.df = None  # DataFrame to store uploaded data
        self.filtered_df = None  # DataFrame to store filtered data
        self.pdf_path = None  # Path to the uploaded PDF
        self.pdf_document = None  # PDF document object
        self.base_name = "Invoice"  # Default base name for exported files
        self.current_zoom = 1.0  # Zoom level for PDF preview
        self.text_boxes = []  # List to store text boxes
        self.box_data = []  # List to store box data
        self.sort_orders = {}  # Dictionary to track column sorting order
        self.docx_template_path = None  # Path to the uploaded DOCX template
        self.image_path = None  # Path to the uploaded image
        self.box_column_map = {}  # Dictionary to map boxes to columns

        # Create the main widget and layout
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout(self.central_widget)

        # Top Bar Layout
        self.top_bar_layout = QHBoxLayout()
        self.layout.addLayout(self.top_bar_layout)

        # Add buttons to the top bar
        self.load_button = QPushButton("Load Data")
        self.top_bar_layout.addWidget(self.load_button)
        self.load_button.clicked.connect(self.load_data)

        self.generate_all_invoices_button = QPushButton("Generate All Invoices")
        self.top_bar_layout.addWidget(self.generate_all_invoices_button)
        self.generate_all_invoices_button.clicked.connect(self.generate_all_invoices)

        self.invoice_button = QPushButton("Create Invoice")
        self.top_bar_layout.addWidget(self.invoice_button)
        self.invoice_button.clicked.connect(self.create_invoice_dialog)

        self.load_image_button = QPushButton("Load Image")
        self.top_bar_layout.addWidget(self.load_image_button)
        self.load_image_button.clicked.connect(self.load_image)

        self.add_box_button = QPushButton("Add Box")
        self.top_bar_layout.addWidget(self.add_box_button)
        self.add_box_button.clicked.connect(self.add_box)

        self.save_structure_button = QPushButton("Save Structure")
        self.top_bar_layout.addWidget(self.save_structure_button)
        self.save_structure_button.clicked.connect(self.save_structure)

        self.load_structure_button = QPushButton("Load Structure")
        self.top_bar_layout.addWidget(self.load_structure_button)
        self.load_structure_button.clicked.connect(self.load_structure)

        self.export_csv_button = QPushButton("Export as CSV")
        self.top_bar_layout.addWidget(self.export_csv_button)
        self.export_csv_button.clicked.connect(lambda: self.export_data("csv"))

        self.export_excel_button = QPushButton("Export as Excel")
        self.top_bar_layout.addWidget(self.export_excel_button)
        self.export_excel_button.clicked.connect(lambda: self.export_data("xlsx"))

        self.export_pdf_button = QPushButton("Export as PDF")
        self.top_bar_layout.addWidget(self.export_pdf_button)
        self.export_pdf_button.clicked.connect(lambda: self.export_data("pdf"))

        # Add a label for instructions
        self.label = QLabel("Load an Excel/CSV file to search, filter, and export data.")
        self.layout.addWidget(self.label)

        # Add search and filter widgets
        self.search_layout = QHBoxLayout()
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Search...")
        self.search_layout.addWidget(self.search_input)

        self.filter_column = QComboBox()
        self.filter_column.addItem("All Columns")
        self.search_layout.addWidget(self.filter_column)

        self.filter_type = QComboBox()
        self.filter_type.addItems(["Contains", "Equals", "Starts with"])
        self.search_layout.addWidget(self.filter_type)

        self.search_button = QPushButton("Search")
        self.search_button.clicked.connect(self.perform_search)
        self.search_layout.addWidget(self.search_button)

        self.layout.addLayout(self.search_layout)

        # Add a table to display the loaded data
        self.table = QTableWidget()
        self.layout.addWidget(self.table)

        # DOCX Upload button
        self.upload_docx_btn = QPushButton("📂 Upload DOCX Template")
        self.upload_docx_btn.clicked.connect(self.upload_template)
        self.top_bar_layout.addWidget(self.upload_docx_btn)

        # Add a button to fill the DOCX template
        self.fill_docx_btn = QPushButton("📝 Fill DOCX Template")
        self.fill_docx_btn.clicked.connect(self.fill_docx_template)
        self.layout.addWidget(self.fill_docx_btn)

        # Initialize PDF preview - REVISED SECTION
        self.pdf_container = QWidget()
        self.pdf_layout = QHBoxLayout(self.pdf_container)

        self.graphics_scene = QGraphicsScene()
        self.graphics_view = QGraphicsView(self.graphics_scene)
        self.pdf_layout.addWidget(self.graphics_view)

        # Add right panel for controls
        self.right_panel = QWidget()
        self.right_layout = QVBoxLayout(self.right_panel)
        self.pdf_layout.addWidget(self.right_panel)

        self.layout.addWidget(self.pdf_container)

    def load_data(self):
        """Loads data from an Excel/CSV file."""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open File", "", "Excel/CSV Files (*.xlsx *.xls *.csv)"
        )

        if not file_path:
            QMessageBox.warning(self, "Error", "No file selected!")
            return

        try:
            if file_path.endswith(".csv"):
                self.df = pd.read_csv(file_path)
            else:
                self.df = pd.read_excel(file_path)

            if self.df.empty:
                QMessageBox.warning(self, "Error", "The file is empty!")
                return

            # Update the filter column dropdown
            self.filter_column.clear()
            self.filter_column.addItem("All Columns")
            self.filter_column.addItems(self.df.columns.tolist())

            # Display the data in the table
            self.display_data_in_table(self.df)

            QMessageBox.information(self, "Success", f"Loaded {len(self.df)} records!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load file: {e}")

    def display_data_in_table(self, data):
        """Displays the loaded data in the table widget."""
        self.table.setRowCount(data.shape[0])
        self.table.setColumnCount(data.shape[1])
        self.table.setHorizontalHeaderLabels(data.columns)

        for i in range(data.shape[0]):
            for j in range(data.shape[1]):
                self.table.setItem(i, j, QTableWidgetItem(str(data.iat[i, j])))

    def perform_search(self):
        """Performs a search on the loaded data."""
        if self.df is None:
            QMessageBox.warning(self, "Error", "No data loaded!")
            return

        search_query = self.search_input.text().strip()
        filter_column = self.filter_column.currentText()
        filter_type = self.filter_type.currentText()

        if not search_query:
            self.display_data_in_table(self.df)
            return

        # Handle comma-separated sub-queries
        sub_queries = [q.strip() for q in search_query.split(',')]
        filtered_data = self.df.copy()

        # Apply queries in priority order
        for q in sub_queries:
            if filter_column == "All Columns":
                filtered_data = filtered_data[
                    filtered_data.apply(lambda row: row.astype(str).str.contains(q, case=False, na=False).any(), axis=1)
                ]
            else:
                if filter_type == "Contains":
                    filtered_data = filtered_data[filtered_data[filter_column].astype(str).str.contains(q, case=False, na=False)]
                elif filter_type == "Equals":
                    filtered_data = filtered_data[filtered_data[filter_column].astype(str) == q]
                elif filter_type == "Starts with":
                    filtered_data = filtered_data[filtered_data[filter_column].astype(str).str.startswith(q, na=False)]

        self.filtered_df = filtered_data

        if self.filtered_df.empty:
            QMessageBox.information(self, "No Results", "No matching records found.")
            return

        self.display_data_in_table(self.filtered_df)

    def export_data(self, format):
        """Exports the filtered data to the specified format."""
        if self.filtered_df is None or self.filtered_df.empty:
            QMessageBox.warning(self, "Error", "No filtered data to export!")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save File", "", f"{format.upper()} Files (*.{format})"
        )

        if not file_path:
            return

        try:
            if format == "csv":
                self.filtered_df.to_csv(file_path, index=False)
            elif format == "xlsx":
                self.filtered_df.to_excel(file_path, index=False)
            elif format == "pdf":
                self.save_df_as_pdf(self.filtered_df, file_path)

            QMessageBox.information(self, "Success", f"Data exported as {format.upper()} successfully!")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to export data: {e}")

    def save_df_as_pdf(self, df, save_path):
        """Saves the DataFrame as a PDF."""
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Add title
        pdf.set_font("Arial", "B", 16)
        pdf.cell(200, 10, "Exported Data", ln=True, align='C')

        # Add table headers
        pdf.set_font("Arial", "B", 12)
        for col in df.columns:
            pdf.cell(40, 10, col, 1)
        pdf.ln()

        # Add table data
        pdf.set_font("Arial", size=12)
        for _, row in df.iterrows():
            for col in df.columns:
                pdf.cell(40, 10, str(row[col]), 1)
            pdf.ln()

        # Save the PDF
        pdf.output(save_path)
        print(f"✅ PDF saved: {save_path}")

    def create_invoice_dialog(self):
        """Opens a dialog to create an invoice."""
        if self.df is None:
            QMessageBox.warning(self, "Error", "No data loaded!")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("Create Invoice")
        layout = QVBoxLayout(dialog)

        # Select Columns
        columns_label = QLabel("Select Columns:")
        layout.addWidget(columns_label)
        columns_list = QListWidget()
        columns_list.setSelectionMode(QListWidget.MultiSelection)
        for col in self.df.columns:
            item = QListWidgetItem(col)
            columns_list.addItem(item)
        layout.addWidget(columns_list)

        # Invoice Details
        details_label = QLabel("Invoice Details:")
        layout.addWidget(details_label)
        form_layout = QFormLayout()
        invoice_number = QLineEdit()
        form_layout.addRow("Invoice Number:", invoice_number)
        customer_name = QLineEdit()
        form_layout.addRow("Customer Name:", customer_name)
        invoice_date = QLineEdit()
        form_layout.addRow("Invoice Date:", invoice_date)
        layout.addLayout(form_layout)

        # Buttons
        button_box = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)

        dialog.exec_()

    def generate_all_invoices(self):
        """Generates a separate invoice for each row in the DataFrame."""
        if self.df is None:
            QMessageBox.warning(self, "Error", "No data loaded!")
            return

        output_folder = QFileDialog.getExistingDirectory(self, "Select Output Folder")
        if not output_folder:
            return

        for index, row in self.df.iterrows():
            invoice_number = str(row.get("Invoice Number", f"INV-{index + 1}"))  # Get invoice number or create one
            customer_name = str(row.get("Customer Name", f"Customer {index + 1}"))  # Get customer name or create one
            invoice_date = str(row.get("Invoice Date", "N/A"))  # Get invoice date or set to N/A

            # Create a temporary DataFrame with only the current row
            temp_df = pd.DataFrame([row])

            # Generate the invoice
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            # Invoice Header
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "Invoice", ln=True, align='C')
            pdf.set_font("Arial", size=12)

            # Customer Information
            address_columns = ["Address", "City", "Zip"]  # Example address columns
            if address_columns:
                pdf.cell(0, 10, "Customer Information:", ln=True)
                address_text = ""
                for col in address_columns:
                    if col in temp_df.columns:
                        address_text += f"{col}: {temp_df.iloc[0][col]}\n"
                pdf.multi_cell(0, 10, address_text)

            # Invoice Details
            pdf.cell(0, 10, f"Invoice Number: {invoice_number}", ln=True)
            pdf.cell(0, 10, f"Customer Name: {customer_name}", ln=True)
            pdf.cell(0, 10, f"Invoice Date: {invoice_date}", ln=True)
            pdf.ln(10)

            # Table Headers
            pdf.set_font("Arial", "B", 12)
            for col in self.df.columns:
                pdf.cell(40, 10, col, 1)
            pdf.ln()

            # Table Data
            pdf.set_font("Arial", size=12)
            for col in self.df.columns:
                pdf.cell(40, 10, str(row[col]), 1)
            pdf.ln()

            # Other Sections
            other_sections = [{"title": "Notes", "column": "Notes"}]  # Example other sections
            if other_sections:
                pdf.ln(10)
                for section in other_sections:
                    if section["column"] in temp_df.columns:
                        pdf.cell(0, 10, f"{section['title']}: {temp_df.iloc[0][section['column']]}", ln=True)

            # Save the PDF
            file_path = os.path.join(output_folder, f"Invoice_{index + 1}.pdf")
            pdf.output(file_path)

        QMessageBox.information(self, "Success", f"{len(self.df)} invoices generated successfully!")

    def load_image(self):
        """Loads an image for the PDF preview."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Open Image", "", "Image Files (*.png *.jpg *.jpeg)")
        if file_path:
            self.image_path = file_path
            self.create_pdf_preview()

    def create_pdf_preview(self):
        """Creates a PDF preview with the loaded image."""
        if self.image_path:
            pixmap = QPixmap(self.image_path)
            self.pdf_label = QLabel()
            self.pdf_label.setPixmap(pixmap)
            self.scroll_area = QScrollArea()
            self.scroll_area.setWidget(self.pdf_label)
            self.graphics_scene = QGraphicsScene()
            self.graphics_view = QGraphicsView(self.graphics_scene)
            self.graphics_view.setGeometry(self.scroll_area.geometry())
            self.graphics_view.setStyleSheet("background: transparent;")
            self.graphics_view.setAttribute(Qt.WA_TranslucentBackground)
            self.layout.addWidget(self.scroll_area)
            self.layout.addWidget(self.graphics_view)
            for rect, dropdown in self.box_column_map.items():
                self.graphics_scene.addItem(rect)
                self.layout.addWidget(dropdown)

    def add_box(self):
        """Adds a resizable box and column selection dropdown."""
        rect = QGraphicsRectItem(100, 100, 200, 50)
        self.graphics_scene.addItem(rect)
        column_dropdown = QComboBox()
        column_dropdown.addItems(self.df.columns.tolist())
        self.layout.addWidget(column_dropdown)
        self.box_column_map[rect] = column_dropdown

    def save_structure(self):
        """Saves the structure of boxes and columns to a JSON file."""
        structure = []
        for rect, dropdown in self.box_column_map.items():
            structure.append({
                "x": rect.rect().x(),
                "y": rect.rect().y(),
                "width": rect.rect().width(),
                "height": rect.rect().height(),
                "column": dropdown.currentText()
            })
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Structure", "", "JSON Files (*.json)")
        if file_path:
            with open(file_path, 'w') as f:
                json.dump(structure, f)

    def load_structure(self):
        """Loads the structure of boxes and columns from a JSON file."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Load Structure", "", "JSON Files (*.json)")
        if file_path:
            with open(file_path, 'r') as f:
                structure = json.load(f)
            self.box_column_map = {}
            self.graphics_scene.clear()
            for item in structure:
                rect = QGraphicsRectItem(item['x'], item['y'], item['width'], item['height'])
                self.graphics_scene.addItem(rect)
                column_dropdown = QComboBox()
                column_dropdown.addItems(self.df.columns.tolist())
                column_dropdown.setCurrentText(item['column'])
                self.layout.addWidget(column_dropdown)
                self.box_column_map[rect] = column_dropdown
            self.create_pdf_preview()

    def generate_pdf_with_boxes(self, output_path):
        """Generates a PDF with boxes and text."""
        doc = fitz.open()  # Create empty PDF
        page = doc.new_page()
        if self.image_path:
            rect = page.rect
            page.insert_image(rect, filename=self.image_path)
        for rect, column_dropdown in self.box_column_map.items():
            column_name = column_dropdown.currentText()
            text = str(self.df.iloc[0][column_name])
            x = rect.rect().x()
            y = rect.rect().y()
            page.insert_text((x, y), text)
        doc.save(output_path)

    def upload_template(self):
        """Handles DOCX template upload with validation"""
        docs_path = QStandardPaths.writableLocation(QStandardPaths.DocumentsLocation)

        try:
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Select DOCX Template",
                docs_path,
                "Word Documents (*.docx);;All Files (*)"
            )

            if not file_path:
                return False

            # Validate file
            file_info = QFileInfo(file_path)

            # Check file exists
            if not file_info.exists():
                QMessageBox.critical(self, "File Not Found", "The selected file does not exist.")
                return False

            # Check file size (max 20MB)
            max_size_mb = 20
            file_size_mb = file_info.size() / (1024 * 1024)
            if file_size_mb > max_size_mb:
                QMessageBox.critical(self, "File Too Large", f"File exceeds maximum size of {max_size_mb}MB")
                return False

            # Check read permissions
            if not file_info.isReadable():
                QMessageBox.critical(self, "Permission Denied", "You don't have permission to read this file.")
                return False

            # Validate DOCX structure
            try:
                from docx import Document
                doc = Document(file_path)
                if not doc.paragraphs and not doc.tables:
                    QMessageBox.warning(self, "Empty Document", "The document appears to be empty or corrupted.")
                    return False

                # Test saving a dummy version (checks for write permissions)
                temp_path = os.path.join(QStandardPaths.writableLocation(
                    QStandardPaths.TempLocation), "temp_validation.docx")
                doc.save(temp_path)
                os.remove(temp_path)

                self.docx_template_path = file_path
                QMessageBox.information(self, "Success", "DOCX template loaded successfully!")
                return True

            except Exception as e:
                QMessageBox.critical(self, "Document Error", f"Failed to process document: {str(e)}")
                return False

        except Exception as e:
            QMessageBox.critical(self, "Unexpected Error", f"An unexpected error occurred:\n{str(e)}")
            return False

    def fill_docx_template(self):
        """Fills the DOCX template with data from the DataFrame."""
        if not hasattr(self, 'docx_template_path') or not self.docx_template_path:
            QMessageBox.critical(self, "Error", "No DOCX template uploaded!")
            return

        if self.df is None or self.df.empty:
            QMessageBox.critical(self, "Error", "No data loaded!")
            return

        output_folder = QFileDialog.getExistingDirectory(self, "Select Output Folder")
        if not output_folder:
            return

        try:
            # Create a DataMapper instance
            mapper = DataMapper(self)

            # Process the documents
            result = mapper.map_data_to_docx(
                self.docx_template_path,
                self.df,
                output_folder
            )

            if result:
                QMessageBox.information(
                    self,
                    "Success",
                    f"Successfully generated {len(result)} documents in:\n{output_folder}"
                )
            else:
                QMessageBox.warning(
                    self,
                    "Warning",
                    "Documents were not generated. Please check the logs."
                )

        except Exception as e:
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to generate documents:\n{str(e)}"
            )
            logging.error(f"Document generation failed: {str(e)}", exc_info=True)

    def validate_docx_template(self):
        """Validates that the DOCX template contains all required placeholders."""
        try:
            doc = Document(self.docx_template_path)
            placeholders_found = set()

            for para in doc.paragraphs:
                for col in self.df.columns:
                    placeholder = f"{{{{{col.strip()}}}}}"
                    if placeholder in para.text:
                        placeholders_found.add(col)

            # Check if all placeholders are found
            missing_placeholders = set(self.df.columns) - placeholders_found
            if missing_placeholders:
                logger.error(f"❌ Missing placeholders in template: {missing_placeholders}")
                return False

            return True

        except Exception as e:
            logger.error(f"❌ Error validating template: {e}")
            return False

    def replace_placeholder_in_runs(self, paragraph, placeholder, replacement):
        """Replaces placeholders in a paragraph while preserving formatting."""
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainApp()
    window.show()
    sys.exit(app.exec_())