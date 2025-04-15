import os
import re
import logging
import pandas as pd
from datetime import datetime
from num2words import num2words
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import ttkbootstrap as tb
import sys
from docx.shared import Pt
from typing import Dict, List, Optional, Set
from docx2pdf import convert
import tempfile

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler("document_generator.log"),
        logging.StreamHandler()
    ]
)


class DocumentGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Generator 3.0")
        self.root.geometry("1200x800")

        # Initialize variables
        self.input_file = None
        self.output_folder = None
        self.current_data = None
        self.templates = {
            "Tax Invoice": None,
            "Credit Note": None,
            "Debit Note": None,
            "Eligible": None,
            "Ineligible": None
        }

        # Setup UI
        self.setup_ui()

        # Load default templates
        self.load_default_templates()

    def setup_ui(self):
        """Setup the main user interface"""
        # Main container
        main_frame = tb.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # Left panel - Controls
        control_frame = tb.Frame(main_frame)
        control_frame.pack(side=tk.LEFT, fill=tk.Y, padx=10, pady=10)

        # Control buttons
        btn_data = tb.Button(control_frame, text="📂 Upload Data File",
                             command=self.upload_data_file)
        btn_data.pack(fill=tk.X, padx=10, pady=5)

        btn_output = tb.Button(control_frame, text="📁 Select Output Folder",
                               command=self.select_output_folder)
        btn_output.pack(fill=tk.X, padx=10, pady=5)

        self.btn_start = tb.Button(
            control_frame,
            text="🚀 Generate DOCUMENT",
            bootstyle="success",
            command=self.start_processing
        )
        self.btn_start.pack(fill=tk.X, padx=10, pady=20)

        # Template status labels
        self.lbl_tax_invoice = tb.Label(control_frame, text="❌ Tax Invoice Template: Not Loaded", bootstyle="danger")
        self.lbl_tax_invoice.pack(fill=tk.X, padx=10, pady=5)

        self.lbl_credit_note = tb.Label(control_frame, text="❌ Credit Note Template: Not Loaded", bootstyle="danger")
        self.lbl_credit_note.pack(fill=tk.X, padx=10, pady=5)

        self.lbl_debit_note = tb.Label(control_frame, text="❌ Debit Note Template: Not Loaded", bootstyle="danger")
        self.lbl_debit_note.pack(fill=tk.X, padx=10, pady=5)

        self.lbl_eligible = tb.Label(control_frame, text="❌ Eligible Template: Not Loaded", bootstyle="danger")
        self.lbl_eligible.pack(fill=tk.X, padx=10, pady=5)

        self.lbl_ineligible = tb.Label(control_frame, text="❌ Ineligible Template: Not Loaded", bootstyle="danger")
        self.lbl_ineligible.pack(fill=tk.X, padx=10, pady=5)

        # Status labels
        self.lbl_data = tb.Label(control_frame, text="No Data File Loaded", bootstyle="secondary")
        self.lbl_data.pack(fill=tk.X, padx=10, pady=5)

        self.lbl_output = tb.Label(control_frame, text="No Output Folder Selected", bootstyle="secondary")
        self.lbl_output.pack(fill=tk.X, padx=10, pady=5)

        # Right panel - Data Preview
        preview_frame = tb.Frame(main_frame)
        preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        preview_label = tb.Label(preview_frame, text="Data Preview", bootstyle="primary")
        preview_label.pack(fill=tk.X, pady=5)

        # Treeview with scrollbars
        self.tree = ttk.Treeview(preview_frame)
        yscroll = ttk.Scrollbar(preview_frame, orient="vertical", command=self.tree.yview)
        xscroll = ttk.Scrollbar(preview_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)

        yscroll.pack(side=tk.RIGHT, fill=tk.Y)
        xscroll.pack(side=tk.BOTTOM, fill=tk.X)
        self.tree.pack(fill=tk.BOTH, expand=True)

    def load_default_templates(self):
        """Load default templates from the templates folder"""
        try:
            # Get the directory where the executable or script is located
            if getattr(sys, 'frozen', False):
                base_path = sys._MEIPASS
                templates_dir = os.path.join(base_path, "templates")
            else:
                base_path = os.path.dirname(os.path.abspath(__file__))
                templates_dir = os.path.join(base_path, "templates")

            # Try to load each template
            template_files = {
                "Tax Invoice": "Tax-Note.docx",
                "Credit Note": "Tax-Note.docx",
                "Debit Note": "Tax-Note.docx",
                "Eligible": "Eligible_template.docx",
                "Ineligible": "Ineligible_template.docx"
            }

            for template_type, filename in template_files.items():
                template_path = os.path.join(templates_dir, filename)
                if os.path.exists(template_path):
                    self.templates[template_type] = template_path
                    logging.info(f"Loaded {template_type} template: {template_path}")
                else:
                    logging.warning(f"Template not found: {template_path}")

            # Update UI labels
            self.update_template_status_labels()

        except Exception as e:
            logging.error(f"Error loading templates: {str(e)}")
            messagebox.showerror("Error", f"Failed to load templates: {str(e)}")

    def update_template_status_labels(self):
        """Update the template status labels based on loaded templates"""
        for template_type, path in self.templates.items():
            label = getattr(self, f"lbl_{template_type.lower().replace(' ', '_')}")
            if path:
                label.config(
                    text=f"✅ {template_type}: {os.path.basename(path)}",
                    bootstyle="success"
                )
            else:
                label.config(
                    text=f"❌ {template_type}: Not Found",
                    bootstyle="danger"
                )

    def upload_data_file(self):
        """Handle data file upload"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Excel files", "*.xlsx;*.xls"), ("CSV files", "*.csv")]
        )

        if file_path:
            self.input_file = file_path
            self.lbl_data.config(text=f"📂 {os.path.basename(file_path)} Loaded")
            logging.info(f"Data file loaded: {file_path}")

            try:
                self.current_data = self.read_data_file(file_path)
                if self.current_data is not None:
                    self.display_data(self.current_data)
                    messagebox.showinfo("Success", "Data file loaded and displayed successfully!")
                else:
                    messagebox.showerror("Error", "Failed to read data file.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to load data: {str(e)}")
                logging.error(f"Data loading error: {str(e)}")

    def read_data_file(self, file_path):
        """Read and clean the data file"""
        try:
            # Read Excel file
            df = pd.read_excel(file_path, engine='openpyxl')

            # Clean column names
            df.columns = [col.strip().upper().replace(' ', '_') for col in df.columns]

            # Fix eligibility column typos
            if 'ELIGIBLE/INELIGIBLE' in df.columns:
                df['ELIGIBLE/INELIGIBLE'] = (
                    df['ELIGIBLE/INELIGIBLE']
                    .str.strip()
                    .str.lower()
                    .replace({
                        'inelgible': 'ineligible',
                        'inellgible': 'ineligible'
                    })
                )

            # Convert numeric columns
            tax_cols = [col for col in df.columns if any(x in col for x in ['IGST', 'CGST', 'SGST', 'UTGST', 'AMOUNT'])]
            for col in tax_cols:
                if col in df.columns:
                    df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

            # Remove empty rows
            df = df.dropna(how='all')

            return df
        except Exception as e:
            logging.error(f"Error reading {file_path}: {str(e)}")
            return None

    def display_data(self, data):
        """Display data in the Treeview"""
        # Clear existing data
        self.tree.delete(*self.tree.get_children())

        # Set up columns
        self.tree["columns"] = list(data.columns)
        self.tree["show"] = "headings"

        # Configure columns
        for col in data.columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=100, stretch=False)

        # Insert data rows
        for _, row in data.iterrows():
            self.tree.insert("", "end", values=list(row))

    def select_output_folder(self):
        """Handle output folder selection"""
        folder = filedialog.askdirectory()
        if folder:
            self.output_folder = folder
            self.lbl_output.config(text=f"📁 Output Folder: {folder}")
            logging.info(f"Output folder selected: {folder}")

    def start_processing(self):
        """Start the document generation process"""
        if not all([self.input_file, self.output_folder]):
            messagebox.showerror("Error", "Please select both data file and output folder!")
            return

        if not any(self.templates.values()):
            messagebox.showerror("Error", "No templates loaded! Cannot generate documents.")
            return

        if self.current_data is None or self.current_data.empty:
            messagebox.showerror("Error", "No valid data to process!")
            return

        try:
            # Create output folders
            os.makedirs(self.output_folder, exist_ok=True)

            # Process each row in the data
            success_count = 0
            for idx, row in self.current_data.iterrows():
                try:
                    # Determine document type
                    doc_type = self.determine_document_type(row)
                    if not doc_type:
                        logging.warning(f"Skipping row {idx} - could not determine document type")
                        continue

                    # Prepare data for template
                    row_data = self.prepare_row_data(row, doc_type)
                    if not row_data:
                        logging.error(f"Failed to prepare data for row {idx}")
                        continue

                    # Generate document
                    if self.generate_document(doc_type, row_data, idx):
                        success_count += 1

                except Exception as e:
                    logging.error(f"Error processing row {idx}: {str(e)}", exc_info=True)
                    continue

            messagebox.showinfo(
                "Complete",
                f"Document generation complete!\n\nSuccessfully generated {success_count} documents."
            )

        except Exception as e:
            logging.error(f"Processing error: {str(e)}", exc_info=True)
            messagebox.showerror("Error", f"Processing failed: {str(e)}")

    def determine_document_type(self, row):
        """Determine the type of document to generate"""
        # First try DOCUMENT_TYPE column
        if 'DOCUMENT_TYPE' in row and pd.notna(row['DOCUMENT_TYPE']):
            doc_type = str(row['DOCUMENT_TYPE']).strip().title()
            if doc_type in ['Tax Invoice', 'Credit Note', 'Debit Note']:
                return doc_type

        # Then try ELIGIBLE/INELIGIBLE column
        if 'ELIGIBLE/INELIGIBLE' in row and pd.notna(row['ELIGIBLE/INELIGIBLE']):
            eligibility = str(row['ELIGIBLE/INELIGIBLE']).strip().lower()
            if eligibility in ['eligible', 'ineligible']:
                return eligibility.title()

        # Fallback to checking tax amounts
        if 'IGST_AS_IGST' in row and float(row['IGST_AS_IGST']) > 0:
            return "Eligible" if 'ELIGIBLE' in str(row.get('ELIGIBLE/INELIGIBLE', '')).upper() else "Ineligible"

        logging.warning(f"Could not determine document type for row: {row.to_dict()}")
        return None

    def prepare_row_data(self, row, doc_type):
        """Prepare the data for document generation"""
        try:
            row_data = {}

            # Normalize row keys (handle spaces, cases, etc.)
            normalized_row = {str(k).strip().upper().replace(' ', '_'): v for k, v in row.items()}

            # Common fields
            common_fields = {
                'INVOICE_NUMBER': ['INVOICE_NUMBER', 'INVOICENUMBER'],
                'INVOICE_DATE': ['INVOICE_DATE', 'INVOICEDATE'],
                'ISD_DISTRIBUTOR_NAME': ['ISD_DISTRIBUTOR_NAME', 'ISDDISTRIBUTORNAME'],
                'ISD_DISTRIBUTOR_ADDRESS': ['ISD_DISTRIBUTOR_ADDRESS', 'ISDDISTRIBUTORADDRESS'],
                'ISD_DISTRIBUTOR_STATE': ['ISD_DISTRIBUTOR_STATE', 'ISDDISTRIBUTORSTATE'],
                'ISD_DISTRIBUTOR_PINCODE': ['ISD_DISTRIBUTOR_PINCODE', 'ISDDISTRIBUTORPINCODE'],
                'ISD_DISTRIBUTOR_STATE_CODE': ['ISD_DISTRIBUTOR_STATE_CODE', 'ISDDISTRIBUTORSTATECODE'],
                'ISD_DISTRIBUTOR_GSTIN': ['ISD_DISTRIBUTOR_GSTIN', 'ISDDISTRIBUTORGSTIN'],
                'CREDIT_RECIPIENT_NAME': ['CREDIT_RECIPIENT_NAME', 'CREDITRECIPIENTNAME'],
                'CREDIT_RECIPIENT_ADDRESS': ['CREDIT_RECIPIENT_ADDRESS', 'CREDITRECIPIENTADDRESS'],
                'CREDIT_RECIPIENT_STATE': ['CREDIT_RECIPIENT_STATE', 'CREDITRECIPIENTSTATE'],
                'CREDIT_RECIPIENT_PINCODE': ['CREDIT_RECIPIENT_PINCODE', 'CREDITRECIPIENTPINCODE'],
                'CREDIT_RECIPIENT_STATE_CODE': ['CREDIT_RECIPIENT_STATE_CODE', 'CREDITRECIPIENTSTATECODE'],
                'CREDIT_RECIPIENT_GSTIN': ['CREDIT_RECIPIENT_GSTIN', 'CREDITRECIPIENTGSTIN'],
                'DOCUMENT_TYPE': ['DOCUMENT_TYPE', 'DOCUMENTTYPE'],
                'SUPPLIER_NAME': ['SUPPLIER_NAME', 'SUPPLIERNAME'],
                'SUPPLIER_ADDRESS': ['SUPPLIER_ADDRESS', 'SUPPLIERADDRESS'],
                'SUPPLIER_PINCODE': ['SUPPLIER_PINCODE', 'SUPPLIERPINCODE'],
                'SUPPLIER_STATE': ['SUPPLIER_STATE', 'SUPPLIERSTATE'],
                'SUPPLIER_STATE_CODE': ['SUPPLIER_STATE_CODE', 'SUPPLIERSTATECODE'],
                'SUPPLIER_GSTIN': ['SUPPLIER_GSTIN', 'SUPPLIERGSTIN'],
                'DOCUMENT_NUMBER': ['DOCUMENT_NUMBER', 'DOCUMENTNUMBER'],
                'DOCUMENT_DATE': ['DOCUMENT_DATE', 'DOCUMENTDATE'],
                'VOUCHER_NO': ['VOUCHER_NO', 'VOUCHERNO'],
                'VOUCHER_DATE': ['VOUCHER_DATE', 'VOUCHERDATE'],
                'RECIPIENT_NAME_BILL_TO': ['RECIPIENT_NAME_BILL_TO', 'RECIPIENTNAMEBILLTO'],
                'RECIPIENT_ADDRESS_BILL_TO': ['RECIPIENT_ADDRESS_BILL_TO', 'RECIPIENTADDRESSBILLTO'],
                'RECIPIENT_PINCODE_BILL_TO': ['RECIPIENT_PINCODE_BILL_TO', 'RECIPIENTPINCODEBILLTO'],
                'RECIPIENT_STATE_NAME_BILL_TO': ['RECIPIENT_STATE_NAME_BILL_TO', 'RECIPIENTSTATENAMEBILLTO'],
                'RECIPIENT_STATE_CODE_BILL_TO': ['RECIPIENT_STATE_CODE_BILL_TO', 'RECIPIENTSTATECODEBILLTO'],
                'RECIPIENT_GSTIN_BILL_TO': ['RECIPIENT_GSTIN_BILL_TO', 'RECIPIENTGSTINBILLTO'],
                'POS': ['POS'],
                'RECIPIENT_NAME_SHIP_TO': ['RECIPIENT_NAME_SHIP_TO', 'RECIPIENTNAMESHIPTO'],
                'RECIPIENT_ADDRESS_SHIP_TO': ['RECIPIENT_ADDRESS_SHIP_TO', 'RECIPIENTADDRESSSHIPTO'],
                'RECIPIENT_PINCODE_SHIP_TO': ['RECIPIENT_PINCODE_SHIP_TO', 'RECIPIENTPINCODESHIPTO'],
                'RECIPIENT_STATE_NAME_SHIP_TO': ['RECIPIENT_STATE_NAME_SHIP_TO', 'RECIPIENTSTATENAMESHIPTO'],
                'RECIPIENT_STATE_CODE_SHIP_TO': ['RECIPIENT_STATE_CODE_SHIP_TO', 'RECIPIENTSTATECODESHIPTO'],
                'RECIPIENT_GSTIN_SHIP_TO': ['RECIPIENT_GSTIN_SHIP_TO', 'RECIPIENTGSTINSHIPTO'],
                'DESCRIPTION_OF_GOODS': ['DESCRIPTION_OF_GOODS', 'DESCRIPTIONOFGOODS'],
                'HSN': ['HSN'],
                'QUANTITY': ['QUANTITY'],
                'UNIT': ['UNIT'],
                'UNIT_PRICE': ['UNIT_PRICE', 'UNITPRICE'],
                'DISCOUNT': ['DISCOUNT'],
                'TAX_RATE': ['TAX_RATE', 'TAXRATE'],
                'BENEFICIARY_NAME': ['BENEFICIARY_NAME', 'BENEFICIARYNAME'],
                'BANK_NAME': ['BANK_NAME', 'BANKNAME'],
                'BANK_ADDRESS': ['BANK_ADDRESS', 'BANKADDRESS'],
                'BANK_ACCOUNT_NO': ['BANK_ACCOUNT_NO', 'BANKACCOUNTNO'],
                'BANK_IFSC_CODE': ['BANK_IFSC_CODE', 'BANKIFSCCODE'],

                'REG_OFFICE': ['REG_OFFICE', 'REGOFFICE'],
                'CIN': ['CIN'],
                'E_MAIL': ['E_MAIL', 'EMAIL'],
                'WEBSITE': ['WEBSITE']
            }

            # Process all fields with their variations
            for standard_name, variations in common_fields.items():
                for variation in variations:
                    if variation in normalized_row and pd.notna(normalized_row[variation]):
                        row_data[standard_name] = self.format_value(normalized_row[variation], standard_name)
                        break
                else:
                    if standard_name in normalized_row and pd.notna(normalized_row[standard_name]):
                        row_data[standard_name] = self.format_value(normalized_row[standard_name], standard_name)

            # Tax fields - handle differently based on document type
            if doc_type in ['Eligible', 'Ineligible']:
                # For ISD documents
                tax_fields = {
                    'IGST_AS_IGST': ['IGST_AS_IGST'],
                    'CGST_AS_IGST': ['CGST_AS_IGST'],
                    'SGST_UTGST_AS_IGST': ['SGST_UTGST_AS_IGST'],
                    'IGST_SUM': ['IGST_SUM'],
                    'CGST_AS_CGST': ['CGST_AS_CGST'],
                    'CGST_SUM': ['CGST_SUM'],
                    'SGST_UTGST_AS_SGST_UTGST': ['SGST_UTGST_AS_SGST_UTGST'],
                    'SGST_UTGST_SUM': ['SGST_UTGST_SUM'],
                    'AMOUNT': ['AMOUNT']
                }
            else:
                # For Tax Invoice/Credit Note/Debit Note
                tax_fields = {
                    'TAXABLE_VALUE': ['TAXABLE_VALUE'],
                    'IGST_SUM': ['IGST_SUM'],
                    'CGST_SUM': ['CGST_SUM'],
                    'SGST_SUM': ['SGST_SUM', 'SGST_UTGST_SUM'],
                    'AMOUNT': ['AMOUNT']
                }

            for standard_name, variations in tax_fields.items():
                for variation in variations:
                    if variation in normalized_row and pd.notna(normalized_row[variation]):
                        row_data[standard_name] = self.format_value(normalized_row[variation], standard_name)
                        break
                else:
                    if standard_name in normalized_row and pd.notna(normalized_row[standard_name]):
                        row_data[standard_name] = self.format_value(normalized_row[standard_name], standard_name)

            # Calculate sums if not provided
            if 'IGST_SUM' not in row_data and all(f in normalized_row and pd.notna(normalized_row[f]) for f in
                                                  ['IGST_AS_IGST', 'CGST_AS_IGST', 'SGST_UTGST_AS_IGST']):
                row_data['IGST_SUM'] = sum([
                    float(normalized_row.get('IGST_AS_IGST', 0)),
                    float(normalized_row.get('CGST_AS_CGST', 0)),
                    float(normalized_row.get('SGST_UTGST_AS_IGST', 0))
                ])

            if 'CGST_SUM' not in row_data and 'CGST_AS_CGST' in normalized_row and pd.notna(
                    normalized_row['CGST_AS_CGST']):
                row_data['CGST_SUM'] = float(normalized_row.get('CGST_AS_CGST', 0))

            if 'SGST_UTGST_SUM' not in row_data and 'SGST_UTGST_AS_SGST_UTGST' in normalized_row and pd.notna(
                    normalized_row['SGST_UTGST_AS_SGST_UTGST']):
                row_data['SGST_UTGST_SUM'] = float(normalized_row.get('SGST_UTGST_AS_SGST_UTGST', 0))

            if 'AMOUNT' not in row_data and all(f in row_data for f in ['IGST_SUM', 'CGST_SUM', 'SGST_UTGST_SUM']):
                row_data['AMOUNT'] = sum([
                    float(row_data.get('IGST_SUM', 0)),
                    float(row_data.get('CGST_SUM', 0)),
                    float(row_data.get('SGST_UTGST_SUM', 0))
                ])

            # Generate amount in words
            try:
                amount_str = str(row_data.get('AMOUNT', '0')).replace(',', '')
                amount = float(amount_str) if amount_str else 0
                if amount % 1 == 0:
                    words = num2words(int(amount), lang='en_IN').title()
                    row_data['amount_in_words'] = f"{words} Rupees Only"
                else:
                    rupees = int(amount)
                    paise = round((amount - rupees) * 100)
                    rupee_words = num2words(rupees, lang='en_IN').title()
                    paise_words = num2words(paise, lang='en_IN').title()
                    row_data['amount_in_words'] = f"{rupee_words} Rupees and {paise_words} Paise Only"
            except Exception as e:
                logging.error(f"Amount conversion error: {str(e)}")
                row_data['amount_in_words'] = "Rupees Only"

            return row_data

        except Exception as e:
            logging.error(f"Error preparing row data: {str(e)}", exc_info=True)
            return None

    def format_value(self, value, key=None):
        """Format values for display in the document"""
        if pd.isna(value) or value in ['', None]:
            return ""

        # Handle numpy types
        if hasattr(value, 'item'):
            value = value.item()

        # Format amounts with 2 decimal places
        if key and any(x in str(key).lower() for x in ['amount', 'igst', 'cgst', 'sgst']):
            try:
                # Remove any existing formatting
                if isinstance(value, str):
                    value = value.replace(',', '').replace('[', '').replace(']', '')
                return "{:,.2f}".format(float(value))
            except:
                return str(value)

        # Special handling for GSTIN (format with spaces)
        if key and 'gstin' in key.lower() and isinstance(value, str) and len(value) == 25:
            return f"{value[0:25]}"

        # Format dates
        if key and 'date' in key.lower():
            try:
                if isinstance(value, str):
                    return value
                return value.strftime('%d-%m-%Y') if hasattr(value, 'strftime') else str(value)
            except:
                return str(value)

        return str(value).strip()

    def generate_document(self, doc_type, row_data, idx):
        """Generate a document and save as PDF, then clean up DOCX"""
        try:
            template_path = self.templates.get(doc_type)
            if not template_path or not os.path.exists(template_path):
                logging.error(f"No template found for document type: {doc_type}")
                return False

            # Load the template
            doc = Document(template_path)

            # Replace placeholders
            if not self.replace_all_placeholders(doc, row_data):
                logging.error(f"Failed to replace placeholders for row {idx}")
                return False

            # Create temp directory for DOCX files
            temp_dir = os.path.join(self.output_folder, "temp_docx")
            os.makedirs(temp_dir, exist_ok=True)

            # Generate filenames
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            invoice_num = str(row_data.get('INVOICE_NUMBER', f"ROW_{idx}")).strip()

            # Temporary DOCX path
            docx_filename = f"{doc_type.replace(' ', '_')}_{invoice_num}_{timestamp}.docx"
            docx_path = os.path.join(temp_dir, docx_filename)

            # Final PDF path
            pdf_filename = f"{doc_type.replace(' ', '_')}_{invoice_num}_{timestamp}.pdf"
            pdf_path = os.path.join(self.output_folder, pdf_filename)

            # Save the DOCX temporarily
            doc.save(docx_path)
            logging.info(f"Temporary DOCX created: {docx_path}")

            try:
                # Convert to PDF
                convert(docx_path, pdf_path)
                logging.info(f"PDF generated: {pdf_path}")

                # Delete the DOCX file
                os.remove(docx_path)
                logging.info(f"Deleted temporary DOCX: {docx_path}")

                # Remove temp directory if empty
                try:
                    os.rmdir(temp_dir)
                except OSError:
                    pass  # Directory not empty

                return True

            except Exception as e:
                logging.error(f"Error converting to PDF: {str(e)}")
                return False

        except Exception as e:
            logging.error(f"Error generating document for row {idx}: {str(e)}", exc_info=True)
            return False

    def replace_all_placeholders(self, doc, replacements):
        """Replace all placeholders in the document with comprehensive handling"""
        try:
            used_placeholders = set()
            missing_placeholders = set()

            # Normalize replacement keys (handle all variations)
            normalized_replacements = {
                self._normalize_placeholder_key(k): str(v) if pd.notna(v) else ""
                for k, v in replacements.items()
            }

            # Process all document components
            self._process_document_components(doc, normalized_replacements, used_placeholders, missing_placeholders)

            # Log diagnostics
            self._log_replacement_stats(used_placeholders, missing_placeholders, normalized_replacements)

            return True

        except Exception as e:
            logging.error(f"Error in replace_all_placeholders: {str(e)}", exc_info=True)
            return False

    def _normalize_placeholder_key(self, key):
        """Normalize placeholder keys to consistent format"""
        if not isinstance(key, str):
            key = str(key)
        return (
            key.strip()
            .upper()
            .replace(' ', '_')
            .replace('-', '_')
            .replace('{', '')
            .replace('}', '')
            .replace('[', '')
            .replace(']', '')
            .replace('(', '')
            .replace(')', '')
        )

    def _process_document_components(self, doc, replacements, used_placeholders, missing_placeholders):
        """Process all components of the document"""
        # Process all paragraphs
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph, replacements, used_placeholders, missing_placeholders)

        # Process all tables and their cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Process cell paragraphs
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph, replacements, used_placeholders, missing_placeholders)

                    # Handle nested tables
                    for nested_table in cell.tables:
                        self._process_document_components(nested_table, replacements, used_placeholders,
                                                          missing_placeholders)

        # Process headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                self._process_paragraph(header, replacements, used_placeholders, missing_placeholders)
            for footer in section.footer.paragraphs:
                self._process_paragraph(footer, replacements, used_placeholders, missing_placeholders)

    def _process_paragraph(self, paragraph, replacements, used_placeholders, missing_placeholders):
        """Process a single paragraph for placeholder replacement"""
        if not paragraph.text:
            return

        original_text = paragraph.text
        new_text = original_text

        # Enhanced pattern to catch all placeholder variants
        placeholder_pattern = r'\{\{([^}]+)\}\}'  # Matches {{PLACEHOLDER}} format
        placeholders = re.findall(placeholder_pattern, original_text)

        for placeholder in placeholders:
            normalized_key = self._normalize_placeholder_key(placeholder)

            if normalized_key in replacements:
                # Get the replacement value (without any braces)
                replacement_value = str(replacements[normalized_key])

                # Replace ALL variants of the placeholder (with any brace style)
                # This ensures we catch any format variations
                new_text = new_text.replace(f"{{{{{placeholder}}}}}", replacement_value)  # {{PLACEHOLDER}}
                new_text = new_text.replace(f"{{{placeholder}}}", replacement_value)  # {PLACEHOLDER}
                new_text = new_text.replace(f"{{ {placeholder} }}", replacement_value)  # { PLACEHOLDER }

                used_placeholders.add(normalized_key)
            else:
                missing_placeholders.add(placeholder)

        if new_text != original_text:
            self._update_paragraph_text(paragraph, new_text)

    def _normalize_placeholder_key(self, key):
        """Normalize placeholder keys to consistent format"""
        if not isinstance(key, str):
            key = str(key)
        return (
            key.strip()
            .upper()
            .replace(' ', '_')
            .replace('-', '_')
            .replace('{', '')
            .replace('}', '')
            .replace('[', '')
            .replace(']', '')
            .replace('(', '')
            .replace(')', '')
        )

    def _update_paragraph_text(self, paragraph, new_text):
        """Update paragraph text while preserving formatting"""
        if not paragraph.runs:
            paragraph.text = new_text
            return

        # Clear all runs except first one
        for run in paragraph.runs[1:]:
            run.text = ""

        # Update first run with new text
        first_run = paragraph.runs[0]
        first_run.text = new_text

        # Preserve basic formatting
        first_run.bold = any(run.bold for run in paragraph.runs)
        first_run.italic = any(run.italic for run in paragraph.runs)
        first_run.underline = any(run.underline for run in paragraph.runs)

    def _log_replacement_stats(self, used_placeholders, missing_placeholders, replacements):
        """Log statistics about placeholder replacement"""
        # Log unused replacements
        unused_replacements = set(replacements.keys()) - used_placeholders
        if unused_replacements:
            logging.warning(f"Unused replacement values: {sorted(unused_replacements)}")

        # Log missing placeholders
        if missing_placeholders:
            logging.warning(f"Missing replacements for placeholders: {sorted(missing_placeholders)}")

        # Log success rate
        total_placeholders = len(used_placeholders) + len(missing_placeholders)
        if total_placeholders > 0:
            success_rate = len(used_placeholders) / total_placeholders * 100
            logging.info(f"Placeholder replacement success: {success_rate:.1f}%")


if __name__ == "__main__":
    root = tb.Window(themename="darkly")
    app = DocumentGeneratorApp(root)
    root.mainloop()