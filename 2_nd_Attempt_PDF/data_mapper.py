import os
import re
import logging
from docx import Document
import pandas as pd
from typing import List, Optional, Set, Dict
from datetime import datetime
from copy import deepcopy
from num2words import num2words
from docx.shared import Pt
from typing import Dict

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# Enhanced column mapping with both eligible and ineligible tax fields
COLUMN_MAPPING = {
    # Invoice fields
    'invoicenumber': 'INVOICE_NUMBER',
    'invoicedate': 'INVOICE_DATE',

    # ISD Distributor fields
    'isddistributorgstin': 'ISD_DISTRIBUTOR_GSTIN',
    'isddistributorname': 'ISD_DISTRIBUTOR_NAME',
    'isddistributoraddress': 'ISD_DISTRIBUTOR_ADDRESS',
    'isddistributorstate': 'ISD_DISTRIBUTOR_STATE',
    'isddistributorpincode': 'ISD_DISTRIBUTOR_PINCODE',
    'isddistributorstatecode': 'ISD_DISTRIBUTOR_STATE_CODE',

    # Credit Recipient fields
    'creditrecipientgstin': 'CREDIT_RECIPIENT_GSTIN',
    'creditrecipientname': 'CREDIT_RECIPIENT_NAME',
    'creditrecipientaddress': 'CREDIT_RECIPIENT_ADDRESS',
    'creditrecipientstate': 'CREDIT_RECIPIENT_STATE',
    'creditrecipientpincode': 'CREDIT_RECIPIENT_PINCODE',
    'creditrecipientstatecode': 'CREDIT_RECIPIENT_STATE_CODE',

    # Tax fields - Handle both eligible and ineligible
    # Eligible Tax fields
    'eligibleigstasigst': 'ELIGIBLE_IGST_AS_IGST',
    'eligiblecgstasigst': 'ELIGIBLE_CGST_AS_IGST',
    'eligiblesgstasigst': 'ELIGIBLE_SGST_AS_IGST',
    'eligibleigstsum': 'ELIGIBLE_IGST_SUM',
    'eligiblecgstascgst': 'ELIGIBLE_CGST_AS_CGST',
    'eligiblecgstsum': 'ELIGIBLE_CGST_SUM',
    'eligiblesgstutgstassgstutgst': 'ELIGIBLE_SGST_UTGST_AS_SGST_UTGST',
    'eligiblesgstutgstassgstutgstsum': 'ELIGIBLE_SGST_UTGST_SUM',
    'eligibleamount': 'ELIGIBLE_AMOUNT',
    # Ineligible Tax fields
    'ineligibleigstasigst': 'INELIGIBLE_IGST_AS_IGST',
    'ineligiblecgstasigst': 'INELIGIBLE_CGST_AS_IGST',
    'ineligiblesgstasigst': 'INELIGIBLE_SGST_AS_IGST',
    'ineligibleigstsum': 'INELIGIBLE_IGST_SUM',
    'ineligiblecgstascgst': 'INELIGIBLE_CGST_AS_CGST',
    'ineligiblecgstsum': 'INELIGIBLE_CGST_SUM',
    'ineligiblesgstutgstassgstutgst': 'INELIGIBLE_SGST_UTGST_AS_SGST_UTGST',
    'ineligiblesgstutgstassgstutgstsum': 'INELIGIBLE_SGST_UTGST_SUM',
    'ineligibleamount': 'INELIGIBLE_AMOUNT',
    'cgst': 'CGST',  # Fallback
    'sgst': 'SGST',  # Fallback
    'utgst': 'UTGST',  # Fallback
    'igst': 'IGST',  # Fallback

    # Amount fields
    'amount': 'AMOUNT',
    'total': 'AMOUNT',

    # Contact fields
    'regoffice': 'REG_OFFICE',
    'cin': 'CIN',
    'e-mail': 'E_MAIL',
    'website': 'WEBSITE',

    # Special fields
    'amount_in_words': 'AMOUNT_IN_WORDS'
}


def normalize_column_names(df: pd.DataFrame) -> pd.DataFrame:
    """Enhanced column name normalization"""
    df.columns = [
        col.strip().upper()
        .replace(' ', '_')
        .replace('-', '_')
        .replace('.', '')
        .replace('ELIGABLE', 'ELIGIBLE')  # Fix common typo
        for col in df.columns
    ]
    return df


def map_data_to_docx(template_path: str, data: pd.DataFrame, output_folder: str,
                    is_eligible: bool = True) -> Optional[List[str]]:
    """
    Main function to generate DOCX files with template selection
    Args:
        template_path: Path to the template file
        data: DataFrame containing the data
        output_folder: Output directory for generated files
        is_eligible: Boolean indicating whether to use eligible template
    """
    try:
        if not validate_inputs(template_path, data, output_folder):
            return None

        os.makedirs(output_folder, exist_ok=True)
        generated_files = []
        template_placeholders = scan_template_placeholders(template_path)

        logging.info(f"Processing {len(data)} rows with {'eligible' if is_eligible else 'ineligible'} template")

        for idx, row in data.iterrows():
            try:
                doc = Document(template_path)
                row_data = prepare_row_data(row, template_placeholders, is_eligible)

                if idx == 0:  # Debug info for first row
                    log_debug_info(row, template_placeholders, row_data)

                if not replace_all_placeholders(doc, row_data):
                    logging.error(f"Skipping row {idx} due to replacement errors")
                    continue

                output_path = generate_output_path(output_folder, row_data, idx, is_eligible)
                doc.save(output_path)
                generated_files.append(output_path)
                logging.info(f"Generated: {os.path.basename(output_path)}")

            except Exception as e:
                logging.error(f"Error processing row {idx}: {str(e)}", exc_info=True)
                continue

        return generated_files if generated_files else None

    except Exception as e:
        logging.error(f"Fatal error in document generation: {str(e)}", exc_info=True)
        return None


def validate_amounts(row_data, prefix):
    """Validate that sums match their components"""
    try:
        # Calculate expected sums
        calc_igst = sum([
            float(row_data.get('IGST_AS_IGST', '0').replace(',', '')),
            float(row_data.get('CGST_AS_IGST', '0').replace(',', '')),
            float(row_data.get('SGST_AS_IGST', '0').replace(',', ''))
        ])
        actual_igst = float(row_data.get('IGST_SUM', '0').replace(',', ''))

        calc_cgst = float(row_data.get('CGST_AS_CGST', '0').replace(',', ''))
        actual_cgst = float(row_data.get('CGST_SUM', '0').replace(',', ''))

        calc_sgst = float(row_data.get('SGST_UTGST_AS_SGST_UTGST', '0').replace(',', ''))
        actual_sgst = float(row_data.get('SGST_UTGST_SUM', '0').replace(',', ''))

        # Check for mismatches
        if not math.isclose(calc_igst, actual_igst, rel_tol=0.01):
            logging.warning(f"IGST_SUM mismatch: Calculated {calc_igst} vs {actual_igst}")

        if not math.isclose(calc_cgst, actual_cgst, rel_tol=0.01):
            logging.warning(f"CGST_SUM mismatch: Calculated {calc_cgst} vs {actual_cgst}")

        if not math.isclose(calc_sgst, actual_sgst, rel_tol=0.01):
            logging.warning(f"SGST_UTGST_SUM mismatch: Calculated {calc_sgst} vs {actual_sgst}")

    except Exception as e:
        logging.error(f"Validation error: {str(e)}")
        return False


def prepare_row_data(row, template_placeholders=None, is_eligible=True):
    prefix = "ELIGIBLE_" if is_eligible else "INELIGIBLE_"
    row_data = {}

    # Process individual tax components
    tax_components = {
        'IGST_AS_IGST': f'{prefix}IGST_AS_IGST',
        'CGST_AS_IGST': f'{prefix}CGST_AS_IGST',
        'SGST_AS_IGST': f'{prefix}SGST_AS_IGST',
        'CGST_AS_CGST': f'{prefix}CGST_AS_CGST',
        'SGST_UTGST_AS_SGST_UTGST': f'{prefix}SGST_UTGST_AS_SGST_UTGST'
    }

    # Process each tax component
    for placeholder, col in tax_components.items():
        if col in row:
            row_data[placeholder] = format_value(row[col], placeholder)
        else:
            row_data[placeholder] = "0.00"

    # Calculate sums correctly
    try:
        # IGST SUM (sum of IGST_AS_IGST, CGST_AS_IGST, SGST_AS_IGST)
        igst_sum = sum([
            float(row.get(f'{prefix}IGST_AS_IGST', 0)),
            float(row.get(f'{prefix}CGST_AS_IGST', 0)),
            float(row.get(f'{prefix}SGST_AS_IGST', 0))
        ])
        row_data['IGST_SUM'] = format_value(igst_sum, 'IGST_SUM')

        # CGST SUM (just CGST_AS_CGST)
        cgst_sum = float(row.get(f'{prefix}CGST_AS_CGST', 0))
        row_data['CGST_SUM'] = format_value(cgst_sum, 'CGST_SUM')

        # SGST/UTGST SUM (just SGST_UTGST_AS_SGST_UTGST)
        sgst_sum = float(row.get(f'{prefix}SGST_UTGST_AS_SGST_UTGST', 0))
        row_data['SGST_UTGST_SUM'] = format_value(sgst_sum, 'SGST_UTGST_SUM')

        # TOTAL AMOUNT (sum of all sums)
        total_amount = igst_sum + cgst_sum + sgst_sum
        row_data['AMOUNT'] = format_value(total_amount, 'AMOUNT')

    except Exception as e:
        logging.error(f"Error calculating sums: {str(e)}")
        row_data['IGST_SUM'] = "0.00"
        row_data['CGST_SUM'] = "0.00"
        row_data['SGST_UTGST_SUM'] = "0.00"
        row_data['AMOUNT'] = "0.00"

    # Calculate total amount if not provided or zero
    if 'AMOUNT' not in row_data or float(row_data['AMOUNT'].replace(',', '')) == 0:
        try:
            total = sum([
                float(row.get(f'{prefix}IGST_SUM', 0)),
                float(row.get(f'{prefix}CGST_SUM', 0)),
                float(row.get(f'{prefix}SGST_UTGST_SUM', 0))
            ])
            row_data['AMOUNT'] = format_value(total, 'AMOUNT')
        except Exception as e:
            logging.error(f"Error calculating total amount: {str(e)}")
            row_data['AMOUNT'] = "0.00"

    # Common fields mapping
    common_fields = {
        # Invoice fields
        'Invoice Number': 'INVOICE_NUMBER',
        'Invoice Date': 'INVOICE_DATE',
        # ISD Distributor fields
        'ISD Distributor Name': 'ISD_DISTRIBUTOR_NAME',
        'ISD Distributor Address': 'ISD_DISTRIBUTOR_ADDRESS',
        'ISD Distributor State': 'ISD_DISTRIBUTOR_STATE',
        'ISD Distributor Pincode': 'ISD_DISTRIBUTOR_PINCODE',
        'ISD Distributor State Code': 'ISD_DISTRIBUTOR_STATE_CODE',
        'ISD Distributor GSTIN': 'ISD_DISTRIBUTOR_GSTIN',
        # Credit Recipient fields
        'Credit Recipient Name': 'CREDIT_RECIPIENT_NAME',
        'Credit Recipient Address': 'CREDIT_RECIPIENT_ADDRESS',
        'Credit Recipient State': 'CREDIT_RECIPIENT_STATE',
        'Credit Recipient Pincode': 'CREDIT_RECIPIENT_PINCODE',
        'Credit Recipient State Code': 'CREDIT_RECIPIENT_STATE_CODE',
        'Credit Recipient GSTIN': 'CREDIT_RECIPIENT_GSTIN',
        # Contact fields
        'Reg. Office': 'REG_OFFICE',
        'CIN': 'CIN',
        'E-Mail': 'E_MAIL',
        'Website': 'WEBSITE',
        # Special fields
        'Amount_In_Words': 'AMOUNT_IN_WORDS'
    }

    # Process common fields
    for placeholder, col in common_fields.items():
        if col in row:
            row_data[placeholder] = format_value(row[col], placeholder)

        # First process all tax components
    for placeholder, col in tax_components.items():
        if col in row:
            row_data[placeholder] = format_value(row[col], placeholder)
        else:
            row_data[placeholder] = "0.00"  # Default value if missing

        # Calculate total amount if not provided or zero
    if 'AMOUNT' not in row_data or float(row_data['AMOUNT'].replace(',', '')) == 0:
        try:
            total = sum([
                float(row.get(f'{prefix}IGST_AS_IGST', 0)),
                float(row.get(f'{prefix}CGST_AS_IGST', 0)),
                float(row.get(f'{prefix}SGST_AS_IGST', 0)),
                float(row.get(f'{prefix}CGST_AS_CGST', 0)),
                float(row.get(f'{prefix}SGST_UTGST_AS_SGST_UTGST', 0))
            ])
            row_data['AMOUNT'] = format_value(total, 'AMOUNT')
        except Exception as e:
            logging.error(f"Error calculating total amount: {str(e)}")
            row_data['AMOUNT'] = "0.00"

    # Generate amount in words if needed
    # In prepare_row_data function
    if any('amount_in_words' in ph.lower() for ph in (template_placeholders or [])):
        try:
            amount_str = row_data.get('AMOUNT', '0').replace(',', '').replace('[', '').replace(']', '')
            amount = float(amount_str)

            if amount % 1 == 0:
                words = num2words(int(amount), lang='en_IN').title()
                row_data['amount_in_words'] = f"{words} Rupees Only"
            else:
                rupees = int(amount)
                paise = round((amount - rupees) * 100)
                rupee_words = num2words(rupees, lang='en_IN').title()
                paise_words = num2words(paise, lang='en_IN').title()
                row_data['amount_in_words'] = (
                    f"{rupee_words} Rupees and "
                    f"{paise_words} Paise Only"
                )
        except Exception as e:
            logging.error(f"Amount conversion error: {str(e)}")
            row_data['amount_in_words'] = "Rupees Only"

    return row_data

def safe_float_conversion(value):
    """Safely convert values to float, handling various edge cases"""
    if pd.isna(value) or value in ['', None]:
        return 0.0
    try:
        return float(value)
    except (ValueError, TypeError):
        return 0.0

def replace_all_placeholders(doc: Document, row_data: Dict[str, str]) -> bool:
    """Replace placeholders throughout document with formatting preservation"""
    try:
        logging.info(f"Available placeholders in row_data: {list(row_data.keys())}")
        logging.info("\n=== Placeholder Replacement ===")
        logging.info(f"Amount value: {row_data.get('AMOUNT', 'MISSING')}")
        logging.info(f"Amount in words: {row_data.get('amount_in_words', 'MISSING')}")

        # Rest of the function remains the same
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, row_data)

        # Process all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, row_data)

        # Process headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        replace_in_paragraph(paragraph, row_data)

            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        replace_in_paragraph(paragraph, row_data)

        return True

    except Exception as e:
        logging.error(f"Error replacing placeholders: {str(e)}", exc_info=True)
        return False


def replace_in_paragraph(paragraph, row_data):
    # First combine all runs
    full_text = ''.join(run.text for run in paragraph.runs)

    # Skip if no replacements needed
    if not any(ph in full_text for ph in row_data):
        return

    # Clear existing content
    paragraph.clear()

    # Split text by placeholders
    parts = re.split(r'(\{\{.+?\}\})', full_text)

    for part in parts:
        if part.startswith('{{') and part.endswith('}}'):
            # This is a placeholder
            ph = part[2:-2].strip()  # Remove braces
            value = str(row_data.get(ph, part))  # Get value or keep original if not found
            run = paragraph.add_run(value)
        else:
            # Regular text
            run = paragraph.add_run(part)

        # Preserve original formatting
        if paragraph.runs and paragraph.runs[0].font.name:
            run.font.name = paragraph.runs[0].font.name
        run.font.size = Pt(10)


def format_value(value, key=None) -> str:
    """Enhanced value formatting with special cases"""
    if pd.isna(value) or value in ['', None]:
        return ""

    # Handle numpy types
    if hasattr(value, 'item'):
        value = value.item()

    # Format amounts with 2 decimal places
    if key and any(x in str(key).lower() for x in ['amount', 'igst', 'cgst', 'sgst']):
        try:
            # Remove any existing formatting
            if isinstance(value, str):
                value = value.replace(',', '').replace('[', '').replace(']', '')
            return "{:,.2f}".format(float(value))
        except:
            return str(value)

    # Special formatting for amounts
    if key and 'amount' in key.lower() and isinstance(value, (int, float)):
        return "{:,.2f}".format(value)

    # Special handling for GSTIN (format with spaces)
    if key and 'gstin' in key.lower() and isinstance(value, str) and len(value) == 15:
        return f"{value[0:25]}"

    return str(value).strip()


def scan_template_placeholders(template_path: str) -> Set[str]:
    """
    Scan a DOCX template and extract all unique placeholder variables
    Args:
        template_path: Path to the template DOCX file
    Returns:
        Set of all unique placeholder names found in the template
    """
    placeholders = set()
    try:
        doc = Document(template_path)

        # Pattern to match {{placeholder}} but ignore **bold** markers
        placeholder_pattern = re.compile(r'\{\{\s*([^{}]+?)\s*\}\}(?!\*)')

        def extract_placeholders(text: str):
            return {match.group(1).strip()
                    for match in placeholder_pattern.finditer(text)}

        # Check all paragraphs in main document
        for paragraph in doc.paragraphs:
            placeholders.update(extract_placeholders(paragraph.text))

        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(extract_placeholders(paragraph.text))

        # Check headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for paragraph in header.paragraphs:
                        placeholders.update(extract_placeholders(paragraph.text))

            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for paragraph in footer.paragraphs:
                        placeholders.update(extract_placeholders(paragraph.text))

        # Check for placeholders in runs (in case they're split across runs)
        for paragraph in doc.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            placeholders.update(extract_placeholders(full_text))

        logging.info(f"Found placeholders in template: {placeholders}")
        return placeholders

    except Exception as e:
        logging.error(f"Error scanning template placeholders: {str(e)}")
        return set()


def generate_output_path(output_folder: str, row_data: dict, idx: int,
                         is_eligible: bool) -> str:
    """Generate output path with type prefix and invoice number"""
    invoice_num = str(row_data.get('INVOICE_NUMBER', idx + 1)).strip()
    prefix = "ELIGIBLE" if is_eligible else "INELIGIBLE"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return os.path.join(output_folder, f"{prefix}_ISD_{invoice_num}_{timestamp}.docx")


def log_debug_info(row, template_placeholders, row_data):
    """Enhanced debug logging with more details"""
    logging.info("\n=== DEBUG INFORMATION ===")
    logging.info(f"Template placeholders: {sorted(template_placeholders)}")
    logging.info(f"Data columns: {sorted(row.index.tolist())}")

    logging.info("\n=== PLACEHOLDER MAPPING ===")
    for ph in sorted(template_placeholders):
        norm_ph = ph.lower().replace(' ', '').replace('.', '').replace('-', '')
        data_key = COLUMN_MAPPING.get(norm_ph, "NO MATCH")
        logging.info(f"Template: {ph:25} → Data: {data_key}")

    logging.info("\n=== MATCHED DATA ===")
    for ph, value in sorted(row_data.items()):
        logging.info(f"{ph:25}: {value}")
    logging.info("=====================")

def validate_template(template_path, required_placeholders):
    doc = Document(template_path)
    found_placeholders = scan_template_placeholders(template_path)
    missing = [ph for ph in required_placeholders if ph not in found_placeholders]
    if missing:
        raise ValueError(f"Missing placeholders in template: {missing}")