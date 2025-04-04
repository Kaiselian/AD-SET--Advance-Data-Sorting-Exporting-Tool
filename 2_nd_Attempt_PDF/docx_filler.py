import os
import re
import logging
from docx import Document
from docx.shared import Pt
from typing import Dict, List, Optional
from datetime import datetime

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# Hardcoded bold elements
BOLD_ELEMENTS = {
    'invoicenumber',
    'invoicedate',
    'Details of ISD Distributor: -',
    'Details of Credit Recipient: -',
    'Name:',
    'Adress:',
    'Pin code:',
    'State Name:',
    'State code:',
    'GSTIN:'
}


def fill_docx_template(template_path: str, output_path: str, replacements: Dict[str, str]) -> bool:
    """Fill template with values and apply hardcoded bold formatting"""
    try:
        doc = Document(template_path)

        # Process all paragraphs
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, replacements)

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph, replacements)

        doc.save(output_path)
        return True
    except Exception as e:
        logging.error(f"Error: {str(e)}")
        return False


def process_paragraph(paragraph, replacements):
    """Process paragraph with hardcoded bold formatting"""
    original_text = paragraph.text
    if not original_text:
        return

    # Clear existing content
    paragraph.clear()

    # Split text into parts that need bold formatting
    parts = re.split(r'(' + '|'.join(map(re.escape, BOLD_ELEMENTS)) + r')', original_text)

    for part in parts:
        if not part:
            continue

        run = paragraph.add_run(part)

        # Apply bold if part matches our hardcoded elements
        if part in BOLD_ELEMENTS:
            run.bold = True

        # Replace placeholders if they exist in this part
        for ph, value in replacements.items():
            if ph in part:
                run.text = run.text.replace(ph, str(value))


def replace_all_placeholders(doc: Document, row_data: Dict[str, str]) -> bool:
    """Replace placeholders throughout document with hardcoded bold elements"""
    try:
        # Process all document components
        components = [
            doc.paragraphs,
            *[cell.paragraphs for table in doc.tables
              for row in table.rows
              for cell in row.cells],
            *[section.header.paragraphs for section in doc.sections],
            *[section.footer.paragraphs for section in doc.sections]
        ]

        for paragraphs in components:
            for paragraph in paragraphs:
                process_paragraph(paragraph, row_data)

        return True
    except Exception as e:
        logging.error(f"Error replacing placeholders: {str(e)}")
        return False


def scan_template_placeholders(template_path: str) -> Set[str]:
    """
    Scan a DOCX template and extract all unique placeholder variables
    Args:
        template_path: Path to the template DOCX file
    Returns:
        Set of all unique placeholder names found in the template
    """
    doc = Document(template_path)
    placeholders = set()
    # Match both {{ }} and {[ ]} styles, and clean the names
    pattern = re.compile(r'\{\{?\s*([^{}]+?)\s*\}?\}')

    def scan_text(text: str):
        return {match.group(1).strip() for match in pattern.finditer(text)}

    # Check all document components
    components = [
        doc.paragraphs,
        *[cell.paragraphs for table in doc.tables
          for row in table.rows
          for cell in row.cells],
        *[section.header.paragraphs for section in doc.sections],
        *[section.footer.paragraphs for section in doc.sections]
    ]

    for paragraphs in components:
        for paragraph in paragraphs:
            placeholders.update(scan_text(paragraph.text))

    return placeholders


def generate_output_filename(row_data: Dict, idx: int, is_eligible: bool) -> str:
    """
    Generate a standardized output filename
    Args:
        row_data: Dictionary containing row data
        idx: Row index
        is_eligible: Whether this is an eligible invoice
    Returns:
        str: Generated filename
    """
    invoice_num = str(row_data.get('INVOICE_NUMBER', idx + 1)).strip()
    prefix = "Eligible" if is_eligible else "Ineligible"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{prefix}_ISD_{invoice_num}_{timestamp}.docx"


def validate_template(template_path: str) -> bool:
    """
    Validate that the template exists and is accessible
    Args:
        template_path: Path to the template file
    Returns:
        bool: True if valid, False otherwise
    """
    try:
        if not os.path.exists(template_path):
            logging.error(f"Template file not found: {template_path}")
            return False
        # Try opening the document to verify it's valid
        Document(template_path)
        return True
    except Exception as e:
        logging.error(f"Invalid template file: {str(e)}")
        return False