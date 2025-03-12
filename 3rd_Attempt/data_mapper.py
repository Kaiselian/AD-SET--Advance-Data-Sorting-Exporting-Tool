import os
from docx import Document
import pandas as pd


def extract_text_with_runs(paragraph):
    """Extracts full text from a paragraph, including formatting differences in runs."""
    return "".join(run.text for run in paragraph.runs)


def normalize_placeholder(text):
    """Cleans up placeholder text to avoid issues with spaces and formatting."""
    return text.strip().lower()


def replace_text_in_paragraphs(paragraphs, row_data):
    """Replace placeholders in paragraphs while keeping formatting intact."""
    placeholders_replaced = set()

    for paragraph in paragraphs:
        full_text = extract_text_with_runs(paragraph)

        for col_name, value in row_data.items():
            placeholder = f"{{{{{normalize_placeholder(col_name)}}}}}"
            if placeholder in normalize_placeholder(full_text):
                print(f"✅ Replacing {placeholder} -> {value}")
                for run in paragraph.runs:
                    if placeholder in normalize_placeholder(run.text):
                        run.text = run.text.replace(placeholder, str(value))
                        placeholders_replaced.add(placeholder)

    return placeholders_replaced


def replace_text_in_tables(tables, row_data):
    """Replace placeholders inside table cells while keeping formatting."""
    placeholders_replaced = set()

    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders_replaced |= replace_text_in_paragraphs(cell.paragraphs, row_data)

    return placeholders_replaced


def map_data_to_docx(template_path, data, output_folder="output_docs"):
    """
    Maps each row of Excel data into a separate DOCX template, saving them individually.
    """
    print("🔹 DEBUG: Mapping Data to Template")
    print(f"📄 Template Path: {template_path}")
    print(f"🟢 Found Columns: {list(data.columns)}")

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    generated_files = []

    for idx, row in data.iterrows():
        try:
            doc = Document(template_path)
        except FileNotFoundError:
            print(f"❌ Error: Template file not found at {template_path}")
            return None

        placeholders_replaced = set()
        row_data = {normalize_placeholder(col): str(row[col]) for col in data.columns}

        # Replace placeholders in document body
        placeholders_replaced |= replace_text_in_paragraphs(doc.paragraphs, row_data)

        # Replace placeholders inside tables
        placeholders_replaced |= replace_text_in_tables(doc.tables, row_data)

        # Replace in headers & footers
        for section in doc.sections:
            placeholders_replaced |= replace_text_in_paragraphs(section.header.paragraphs, row_data)
            placeholders_replaced |= replace_text_in_paragraphs(section.footer.paragraphs, row_data)

        # Save each filled document separately
        output_filename = os.path.join(output_folder, f"filled_{idx + 1}.docx")
        try:
            doc.save(output_filename)
            generated_files.append(output_filename)
            print(f"✅ Successfully saved: {output_filename}")
        except Exception as e:
            print(f"❌ Error: Could not save filled DOCX: {e}")

        if not placeholders_replaced:
            print(f"⚠️ WARNING: No placeholders replaced in {output_filename}. Check DOCX formatting.")

    return generated_files
