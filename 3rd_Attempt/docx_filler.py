from docx import Document
import pandas as pd
import os
from docx2pdf import convert


def fill_docx_template(template_path, data, output_folder):
    """Fills a DOCX template with multiple rows from Excel and converts to PDFs."""

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    filled_files = []

    for idx, row in data.iterrows():
        doc = Document(template_path)
        placeholders_replaced = set()

        # Replace text in paragraphs (Preserving Formatting)
        for para in doc.paragraphs:
            for run in para.runs:
                for col in data.columns:
                    placeholder = f"{{{{{col.strip()}}}}}"
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(row[col]))
                        placeholders_replaced.add(placeholder)

        # Replace text in tables
        for table in doc.tables:
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for col in data.columns:
                                placeholder = f"{{{{{col.strip()}}}}}"
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(row[col]))
                                    placeholders_replaced.add(placeholder)

        # Save DOCX
        output_docx = os.path.join(output_folder, f"filled_{idx + 1}.docx")
        doc.save(output_docx)
        filled_files.append(output_docx)

        print(f"✅ DOCX filled: {output_docx}")

        # Convert to PDF
        output_pdf = output_docx.replace(".docx", ".pdf")
        convert(output_docx, output_pdf)

        print(f"📄 Converted to PDF: {output_pdf}")

    print(f"\n✅ All {len(filled_files)} documents processed successfully!")
    return filled_files
