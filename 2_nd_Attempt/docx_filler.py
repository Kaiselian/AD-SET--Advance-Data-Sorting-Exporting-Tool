from docx import Document
import pandas as pd
from docx2pdf import convert
import os

def fill_docx_with_data(template_path, data_path, output_folder):
    # Your function logic here
    pass  # Placeholder

def fill_docx_template(template_path, df, mapping, output_folder):
    """
    Fills a DOCX template with data from a DataFrame and exports individual PDFs.

    :param template_path: Path to the DOCX template file
    :param df: Pandas DataFrame containing data
    :param mapping: Dictionary mapping placeholders to column names
    :param output_folder: Folder to save the generated files
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)  # Ensure output directory exists

    for index, row in df.iterrows():
        try:
            doc = Document(template_path)  # Load the template

            # Replace placeholders in paragraphs
            for para in doc.paragraphs:
                for placeholder, column in mapping.items():
                    if column in df.columns:
                        value = str(row[column]) if pd.notna(row[column]) else ""
                        para.text = para.text.replace(f"{{{placeholder}}}", value)

            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for placeholder, column in mapping.items():
                            if column in df.columns:
                                value = str(row[column]) if pd.notna(row[column]) else ""
                                cell.text = cell.text.replace(f"{{{placeholder}}}", value)

            # Save filled DOCX
            docx_filename = os.path.join(output_folder, f"Filled_{index + 1}.docx")
            doc.save(docx_filename)

            # Convert to PDF
            pdf_filename = os.path.join(output_folder, f"Filled_{index + 1}.pdf")
            convert(docx_filename, pdf_filename)

            print(f"✅ Generated: {pdf_filename}")

        except Exception as e:
            print(f"❌ Error processing row {index + 1}: {e}")

